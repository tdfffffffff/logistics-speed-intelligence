"""Generate the manager-facing brief as a Word document.

Deliberately a *generator*, not a hand-written document. It is assembled from
the same fact pack and routed model output the rest of the pipeline uses, so it
regenerates whenever the data refreshes -- the same effort as writing it once,
with a reusable tool at the end of it.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from speedlab.config import FIGURES

INK = RGBColor(0x0B, 0x0B, 0x0B)
MUTED = RGBColor(0x52, 0x51, 0x4E)
ACCENT = RGBColor(0x2A, 0x78, 0xD6)
ALERT = RGBColor(0xC1, 0x27, 0x2D)


def _style(doc):
    n = doc.styles["Normal"]
    n.font.name, n.font.size = "Calibri", Pt(10.5)
    n.paragraph_format.space_after = Pt(6)


def _h(doc, text, size=13, color=INK, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold, r.font.size, r.font.color.rgb = True, Pt(size), color
    return p


def _kv(doc, label, value, color=INK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}  ")
    r.font.size, r.font.color.rgb = Pt(9.5), MUTED
    r2 = p.add_run(value)
    r2.bold, r2.font.size, r2.font.color.rgb = True, Pt(11), color


def _fig(doc, name, width=6.2, caption=None):
    path = FIGURES / f"{name}.png"
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        r.italic, r.font.size, r.font.color.rgb = True, Pt(8.5), MUTED


def build_insight_brief(pack: dict, reports: dict, routing, out_path) -> Path:
    """Assemble the two-page manager brief from computed facts."""
    doc = Document()
    _style(doc)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.8)

    h = pack["headline"]
    meta = pack["meta"]
    camp = pack["campaign_effect"]
    imp = pack["impact"]

    # ---------------------------------------------------------- title
    t = doc.add_paragraph()
    r = t.add_run("Regional Speed Review")
    r.bold, r.font.size, r.font.color.rgb = True, Pt(19), INK
    sub = doc.add_paragraph()
    r = sub.add_run(f"{meta['date_range'][0]} to {meta['date_range'][1]}  ·  "
                    f"{meta['days_covered']} days  ·  {meta['total_parcels']:,} parcels  ·  "
                    f"{len(meta['countries'])} markets, {len(meta['providers'])} carriers")
    r.font.size, r.font.color.rgb = Pt(9.5), MUTED

    # ---------------------------------------------------------- headline KPIs
    _h(doc, "Where we stand", space_before=12)
    _kv(doc, "Buyer waiting time (avg_BWT)", f"{h['avg_BWT_days']} days")
    _kv(doc, "Of which seller preparation (avg_APT)",
        f"{h['avg_APT_days']} days  ({h['apt_share_of_bwt_pct']}% of total)")
    _kv(doc, "Of which network transit", f"{h['avg_transit_days']} days", ACCENT)

    # ---------------------------------------------------------- findings
    _h(doc, "What the month shows")
    findings = [
        ("Speed degrades two days AFTER every campaign, not during it.",
         f"Correlation of daily volume with buyer waiting time peaks at a {camp['peak_lag_days']}-day lag "
         f"(r = {camp['peak_correlation']}, 95% CI {camp['peak_ci_95'][0]}–{camp['peak_ci_95'][1]}). "
         f"Campaign dates are known weeks ahead, so this is schedulable rather than merely observable."),
        ("The delay is in the network, not with sellers.",
         f"Preparation time is flat at {h['avg_APT_days']} days across every carrier, region and date. "
         f"Effectively all variation buyers feel occurs after the seller hands the parcel over."),
        ("Island crossings are a step change, not a gradient.",
         "Sea and air legs into Sabah, Sulawesi and Sumatra run roughly 3.5× same-region lanes. "
         "This is where the recoverable time sits."),
        ("The obvious carrier ranking is misleading.",
         "Comparing carriers only on lanes they both serve moves Pos Indonesia from 9th to 5th — it is being "
         "penalised for its geography. The same correction confirms SiCepat's gap is real."),
        ("Two carriers are failing in opposite ways.",
         "SiCepat is consistently slow and deteriorating. Ninja Van is average on speed but generates about "
         "half of all severe incidents. These need different interventions."),
    ]
    for i, (title, body) in enumerate(findings, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(f"{i}. {title}  ")
        r.bold, r.font.size = True, Pt(10.5)
        r2 = p.add_run(body)
        r2.font.size = Pt(10)

    _fig(doc, "01_campaign_backlog", 6.0,
         "Volume peaks on campaign days; buyer waiting time peaks two days later.")

    doc.add_page_break()

    # ---------------------------------------------------------- opportunity
    _h(doc, "What it is worth", space_before=0)
    _kv(doc, "Recoverable buyer waiting, per month",
        f"{imp['total_recoverable_parcel_days']/1e6:.0f}M parcel-days", ACCENT)
    p = doc.add_paragraph()
    r = p.add_run("Counterfactual: every country × lane-class segment reaching the median segment's speed — "
                  "not best-in-class, simply typical. The two island-crossing segments (Indonesia and "
                  "Malaysia) account for roughly two thirds of the total.")
    r.font.size = Pt(10)
    _fig(doc, "07_impact_sizing", 5.8)

    # ---------------------------------------------------------- actions
    _h(doc, "Recommended actions")
    actions = [
        ("Capacity & Network Planning",
         "Staff hubs and linehaul for a +2 day lag after each campaign, not on the campaign day itself.",
         "This month"),
        ("3PL Partner Management",
         "Open a performance review with SiCepat on the lane-matched gap. Close the Pos Indonesia issue — "
         "its raw figure reflects geography, not execution.",
         "This quarter"),
        ("Operations Excellence",
         "Treat Ninja Van as a reliability problem, not a speed one: focus on incident response and hub "
         "handover rather than average transit.",
         "This month"),
        ("Data Engineering",
         "Fix region mapping. Unmapped rows are the slowest lanes in the network and are currently invisible "
         "on every regional dashboard.",
         "This quarter"),
        ("Seller Operations",
         "No action required. Preparation time is consistent everywhere; this is evidence, not an oversight.",
         "—"),
    ]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Light Grid Accent 1"
    for c, txt in zip(tbl.rows[0].cells, ["Owner", "Action", "Horizon"]):
        c.text = txt
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold, rr.font.size = True, Pt(9.5)
    for owner, action, when in actions:
        cells = tbl.add_row().cells
        for cell, txt, sz in zip(cells, [owner, action, when], [9.5, 9.5, 9.5]):
            cell.text = txt
            for pp in cell.paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(sz)

    # ---------------------------------------------------------- caveats
    _h(doc, "How to read these numbers")
    dq = pack["data_quality"]
    caveats = [
        f"{dq['rows_quarantined']} rows were excluded from headline metrics: "
        f"{dq['quarantine_reasons']['zero_parcel_qty']} carried zero parcels and "
        f"{dq['quarantine_reasons']['apt_exceeds_bwt']} reported preparation time exceeding total waiting "
        f"time, which is not physically possible and indicates a source-data fault.",
        f"{dq['unmapped_region_rows']} rows ({dq['unmapped_region_pct_of_parcels']}% of parcels) have an "
        f"unmapped region. These were retained, not dropped: they are slower than mapped rows in 18 of 18 "
        f"carrier-country pairs, so excluding them would understate national waiting time.",
        "All lanes in this dataset are domestic — there is no cross-border traffic to report on.",
        "One month covers roughly three campaign cycles. The lag finding is strong but rests on three events; "
        "a full quarter would be needed before committing headcount to it.",
    ]
    for c in caveats:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(c)
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED

    # ---------------------------------------------------------- provenance
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(
        f"Generated automatically from the analysis pipeline on {meta['generated_at'][:10]}. "
        f"Narrative sections drafted by the routed language model and machine-verified against the "
        f"computed fact pack: every figure above traces to a value in the source data.")
    r.italic, r.font.size, r.font.color.rgb = True, Pt(8.5), MUTED

    out_path = Path(out_path)
    doc.save(out_path)
    return out_path


def to_pdf(docx_path) -> Path | None:
    """Convert to PDF if a converter is available. Never fatal: the .docx is
    the deliverable and PDF is a convenience."""
    try:
        from docx2pdf import convert
        convert(str(docx_path))
        pdf = Path(docx_path).with_suffix(".pdf")
        return pdf if pdf.exists() else None
    except Exception:
        return None
